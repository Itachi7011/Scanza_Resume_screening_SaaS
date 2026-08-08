"""
File -> plaintext extraction. This is the entry point for the whole
pipeline's power — if text extraction is weak, nothing downstream can
recover. Four engines, tried in order of speed/quality tradeoff:

  1. pdfplumber (layout=True + column-aware word clustering) — best
     reading-order preservation for multi-column resumes
  2. PyMuPDF (fitz) — fast, sometimes recovers text pdfplumber misses
     (different font-encoding edge cases)
  3. pdfminer.six (low-level) — a third, independent parsing engine,
     catches PDFs the other two mis-render
  4. OCR (pytesseract + pdf2image) — the real power move: if the first
     three all yield near-empty text, the PDF is almost certainly a
     scanned image with no text layer. We rasterize each page and OCR
     it rather than giving up. This alone handles a huge class of
     resumes (scanned, exported-from-image-editor, old scanned CVs)
     that a naive parser would completely fail on.

We also detect and correctly handle multi-column layouts: pdfplumber's
word-level bounding boxes let us cluster words into columns by x-position
and read column-by-column, top-to-bottom, instead of interleaving lines
from two columns into garbage (the single most common cause of mangled
resume text extraction).
"""
import io
import logging
from typing import List, Tuple

import pdfplumber
import pymupdf as fitz  # PyMuPDF — `import fitz` still works but is deprecated as of PyMuPDF 1.28+
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text

logger = logging.getLogger("resume_worker.text_extraction")

MIN_ACCEPTABLE_CHARS = 60
OCR_DPI = 300


def _cluster_columns(words: list, page_width: float, gap_threshold_ratio: float = 0.04) -> List[List[dict]]:
    """
    Groups words into left-to-right column clusters based on x-position
    gaps, so a two-column resume reads as "full left column, then full
    right column" instead of interleaving both columns line-by-line
    (which produces nonsensical text like "John SoftwareDoe Engineer").
    """
    if not words:
        return []

    sorted_by_x = sorted(words, key=lambda w: w["x0"])
    gap_threshold = page_width * gap_threshold_ratio * 3  # generous gap to avoid over-splitting normal word spacing

    columns: List[List[dict]] = [[sorted_by_x[0]]]
    for w in sorted_by_x[1:]:
        prev_max_x1 = max(x["x1"] for x in columns[-1])
        if w["x0"] - prev_max_x1 > gap_threshold and len(columns[-1]) > 3:
            columns.append([w])
        else:
            columns[-1].append(w)

    # A resume with one wide column will produce one cluster — that's correct,
    # single-column is the common case and this doesn't break it.
    return columns


def _words_to_reading_order_text(page) -> str:
    words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
    if not words:
        return ""

    columns = _cluster_columns(words, page.width)
    lines_out: List[str] = []
    for col in columns:
        col_sorted = sorted(col, key=lambda w: (round(w["top"] / 3), w["x0"]))
        current_top_band = None
        current_line: List[str] = []
        for w in col_sorted:
            band = round(w["top"] / 3)
            if current_top_band is not None and band != current_top_band:
                lines_out.append(" ".join(current_line))
                current_line = []
            current_line.append(w["text"])
            current_top_band = band
        if current_line:
            lines_out.append(" ".join(current_line))
    return "\n".join(lines_out)


def _extract_with_pdfplumber(file_bytes: bytes) -> str:
    text_parts: List[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            layout_text = page.extract_text(layout=True) or ""
            column_text = _words_to_reading_order_text(page)
            text_parts.append(column_text if len(column_text) > len(layout_text) else layout_text)
    return "\n".join(text_parts).strip()


def _extract_with_pymupdf(file_bytes: bytes) -> str:
    text_parts: List[str] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        text_parts.append(page.get_text("text"))
    doc.close()
    return "\n".join(text_parts).strip()


def _extract_with_pdfminer(file_bytes: bytes) -> str:
    return (pdfminer_extract_text(io.BytesIO(file_bytes)) or "").strip()


def _extract_with_ocr(file_bytes: bytes) -> str:
    """Last-resort engine for scanned/image-only PDFs. Rasterizes each page
    at high DPI and runs Tesseract OCR — meaningfully slower than the other
    engines but turns a 0%-extractable scanned resume into a usable one."""
    from pdf2image import convert_from_bytes
    import pytesseract

    images = convert_from_bytes(file_bytes, dpi=OCR_DPI)
    text_parts = [pytesseract.image_to_string(img) for img in images]
    return "\n".join(text_parts).strip()


def _extract_pdf(file_bytes: bytes) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    candidates: List[Tuple[str, str]] = []

    for name, fn in [("pdfplumber", _extract_with_pdfplumber), ("PyMuPDF", _extract_with_pymupdf), ("pdfminer.six", _extract_with_pdfminer)]:
        try:
            text = fn(file_bytes)
            candidates.append((name, text))
        except Exception as e:  # noqa: BLE001 — intentional broad catch, this is a fallback chain
            warnings.append(f"{name} failed: {e}")

    best_name, best_text = max(candidates, key=lambda c: len(c[1])) if candidates else ("none", "")

    if len(best_text) < MIN_ACCEPTABLE_CHARS:
        warnings.append(
            f"Only {len(best_text)} chars extracted via {best_name} — this looks like a scanned/image-only "
            f"PDF. Falling back to OCR (this may take a few extra seconds)."
        )
        try:
            ocr_text = _extract_with_ocr(file_bytes)
            if len(ocr_text) > len(best_text):
                logger.info("OCR fallback recovered %d chars (vs %d from text layer)", len(ocr_text), len(best_text))
                return ocr_text, warnings + ["Text was recovered via OCR — accuracy may be slightly lower than a native text PDF."]
        except Exception as e:  # noqa: BLE001
            warnings.append(f"OCR fallback also failed: {e}")

    return best_text, warnings


def _extract_docx(file_bytes: bytes) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(p for p in parts if p.strip())
        if len(text) < MIN_ACCEPTABLE_CHARS:
            warnings.append("Very little text found in this DOCX — it may rely on text boxes, which aren't fully supported.")
    except Exception as e:  # noqa: BLE001
        warnings.append(f"DOCX parsing failed: {e}")
        text = ""
    return text, warnings


def extract_text(file_bytes: bytes, content_type: str, filename: str) -> Tuple[str, List[str]]:
    """Returns (plaintext, warnings). Dispatches by content type / extension."""
    lower_name = filename.lower()

    if "pdf" in content_type or lower_name.endswith(".pdf"):
        return _extract_pdf(file_bytes)

    if "wordprocessingml" in content_type or lower_name.endswith(".docx"):
        return _extract_docx(file_bytes)

    if lower_name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore"), []

    return "", [f"Unsupported file type: {content_type or lower_name}"]
